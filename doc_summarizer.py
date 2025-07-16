import os
import tempfile
import streamlit as st
from pypdf import PdfReader
from transformers import pipeline
from reportlab.pdfgen import canvas
import docx
import io

os.environ['CURL_CA_BUNDLE'] = ''  # Disable SSL certificate verification for Hugging Face requests

st.set_page_config(page_title="Document Summarizer & Reporter", layout="centered")
st.title("Document Summarizer & Reporter App")

st.markdown("""
Upload a PDF, text, or Word file and get a summary using Hugging Face (facebook/bart-large-cnn).
You can also generate a sample PDF for testing. The app will also generate a basic report for your document.
""")

# Generate sample files

def generate_sample_pdf():
    temp_dir = tempfile.gettempdir()
    sample_path = os.path.join(temp_dir, "sample.pdf")
    c = canvas.Canvas(sample_path)
    c.drawString(100, 750, "This is a sample PDF file for testing PDF summarization.")
    c.drawString(100, 730, "You can upload your own PDF, TXT, or DOCX to get a summary and report.")
    c.save()
    return sample_path

def generate_sample_txt():
    sample_txt = "This is a sample TXT file for testing text summarization.\nYou can upload your own PDF, TXT, or DOCX to get a summary and report."
    return io.BytesIO(sample_txt.encode("utf-8"))

def generate_sample_docx():
    doc_stream = io.BytesIO()
    doc = docx.Document()
    doc.add_paragraph("This is a sample DOCX file for testing Word summarization.")
    doc.add_paragraph("You can upload your own PDF, TXT, or DOCX to get a summary and report.")
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

uploaded_file = st.file_uploader("Upload your document (PDF, TXT, DOCX)", type=["pdf", "txt", "docx"])
doc_text = ""
report = {}

if uploaded_file:
    file_type = uploaded_file.name.split('.')[-1].lower()
    if file_type == "pdf":
        reader = PdfReader(uploaded_file)
        doc_text = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
        report['Type'] = "PDF"
        report['Pages'] = len(reader.pages)
        report['Characters'] = len(doc_text)
        report['Words'] = len(doc_text.split())
    elif file_type == "txt":
        doc_text = uploaded_file.read().decode("utf-8")
        report['Type'] = "Text File"
        report['Characters'] = len(doc_text)
        report['Words'] = len(doc_text.split())
    elif file_type == "docx":
        doc = docx.Document(uploaded_file)
        doc_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        report['Type'] = "Word File"
        report['Paragraphs'] = len(doc.paragraphs)
        report['Characters'] = len(doc_text)
        report['Words'] = len(doc_text.split())
    else:
        st.error("Unsupported file type.")

    st.text_area("Extracted Document Text (first 1000 chars)", doc_text[:1000], height=200)

    # Show basic report
    st.subheader("Document Report")
    for k, v in report.items():
        st.write(f"**{k}:** {v}")

    if st.button("Summarize Document"):
        with st.spinner("Loading summarization model and generating summary..."):
            model_name = "facebook/bart-large-cnn"
            summarizer = pipeline("summarization", model=model_name)
            truncated_text = doc_text[:1000]
            summary = summarizer(truncated_text, max_length=1000, min_length=500, do_sample=False)[0]['summary_text']
        st.subheader("Summary")
        st.write(summary)
else:
    st.info("Upload a PDF, TXT, or DOCX file to get started or generate a sample PDF above.")

# --- Sidebar: Sample Files ---
with st.sidebar:
    st.header("📁 Download Sample Files")
    
    # PDF sample - always available
    sample_path = generate_sample_pdf()
    with open(sample_path, "rb") as f:
        st.download_button("📄 Download Sample PDF", f, file_name="sample.pdf", key="pdf_download")
    
    # TXT sample - always available
    sample_txt = generate_sample_txt()
    st.download_button("📝 Download Sample TXT", sample_txt, file_name="sample.txt", key="txt_download")
    
    # DOCX sample - always available
    sample_docx = generate_sample_docx()
    st.download_button("📋 Download Sample DOCX", sample_docx, file_name="sample.docx", key="docx_download")
